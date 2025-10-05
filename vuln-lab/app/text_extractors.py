"""
Text extraction utilities for different file formats.
Supports PDF, DOC, DOCX, TXT, and other common resume formats.
"""

import os
import logging
from typing import Optional, Union, Dict, Any

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TextExtractor:
    """Main text extraction class that handles different file formats."""
    
    def __init__(self):
        """Initialize the text extractor with available libraries."""
        self.available_extractors = self._check_available_extractors()
    
    def _check_available_extractors(self) -> Dict[str, bool]:
        """Check which extraction libraries are available."""
        available = {
            'pypdf2': False,
            'pymupdf': False,
            'python_docx': False,
            'python_doc': False,
            'textract': False
        }
        
        try:
            import PyPDF2
            available['pypdf2'] = True
        except ImportError:
            pass
        
        try:
            import fitz  # PyMuPDF
            available['pymupdf'] = True
        except ImportError:
            pass
        
        try:
            from docx import Document
            available['python_docx'] = True
        except ImportError:
            pass
        
        try:
            import python_docx2txt
            available['python_doc'] = True
        except ImportError:
            pass
        
        try:
            import textract
            available['textract'] = True
        except ImportError:
            pass
        
        logger.info(f"Available extractors: {available}")
        return available
    
    def extract_text(self, file_path: str = None, file_content: bytes = None, filename: str = None) -> str:
        """
        Extract text from a file based on its extension.
        
        Args:
            file_path: Path to the file (if reading from disk)
            file_content: File content as bytes (if reading from memory)
            filename: Original filename to determine extension
        
        Returns:
            Extracted text content
        """
        if not filename and not file_path:
            raise ValueError("Either filename or file_path must be provided")
        
        # Determine filename and extension
        if not filename:
            filename = os.path.basename(file_path)
        
        extension = os.path.splitext(filename)[1].lower()
        
        try:
            if extension == '.txt':
                return self._extract_txt(file_path, file_content)
            elif extension == '.pdf':
                return self._extract_pdf(file_path, file_content)
            elif extension in ['.docx']:
                return self._extract_docx(file_path, file_content)
            elif extension in ['.doc']:
                return self._extract_doc(file_path, file_content)
            elif extension in ['.rtf']:
                return self._extract_rtf(file_path, file_content)
            else:
                return self._extract_fallback(file_path, file_content, filename)
        
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            return f"[Error extracting text from {filename}: {str(e)}]"
    
    def _extract_txt(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from TXT files."""
        if file_content:
            return file_content.decode('utf-8', errors='ignore')
        elif file_path:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        return ""
    
    def _extract_pdf(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from PDF files using available libraries."""
        text = ""
        
        # Try PyMuPDF first (better quality)
        if self.available_extractors['pymupdf']:
            try:
                import fitz
                if file_content:
                    doc = fitz.open(stream=file_content, filetype="pdf")
                else:
                    doc = fitz.open(file_path)
                
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}")
        
        # Try PyPDF2 as fallback
        if self.available_extractors['pypdf2']:
            try:
                import PyPDF2
                import io
                
                if file_content:
                    pdf_file = io.BytesIO(file_content)
                else:
                    pdf_file = open(file_path, 'rb')
                
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
                
                if not file_content:
                    pdf_file.close()
                
                return text
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")
        
        # Try textract as last resort
        if self.available_extractors['textract'] and file_path:
            try:
                import textract
                text = textract.process(file_path).decode('utf-8')
                return text
            except Exception as e:
                logger.warning(f"Textract failed: {e}")
        
        return f"[PDF text extraction not available - install PyMuPDF, PyPDF2, or textract]"
    
    def _extract_docx(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from DOCX files."""
        if self.available_extractors['python_docx']:
            try:
                from docx import Document
                import io
                
                if file_content:
                    doc = Document(io.BytesIO(file_content))
                else:
                    doc = Document(file_path)
                
                text = []
                print("Extracting text from DOCX...")
                print(doc)
                for paragraph in doc.paragraphs:
                    print(paragraph.get_text())
                    text.append(paragraph.text)
                
                return '\n'.join(text)
            except Exception as e:
                logger.warning(f"python-docx failed: {e}")
        
        # Try textract as fallback
        if self.available_extractors['textract'] and file_path:
            try:
                import textract
                text = textract.process(file_path).decode('utf-8')
                return text
            except Exception as e:
                logger.warning(f"Textract failed: {e}")
        
        return f"[DOCX text extraction not available - install python-docx or textract]"
    
    def _extract_doc(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from DOC files."""
        # Try textract first (best for .doc files)
        if self.available_extractors['textract'] and file_path:
            try:
                import textract
                text = textract.process(file_path).decode('utf-8')
                return text
            except Exception as e:
                logger.warning(f"Textract failed for DOC: {e}")
        
        return f"[DOC text extraction not available - install textract or convert to DOCX]"
    
    def _extract_rtf(self, file_path: str = None, file_content: bytes = None) -> str:
        """Extract text from RTF files."""
        # Try textract
        if self.available_extractors['textract'] and file_path:
            try:
                import textract
                text = textract.process(file_path).decode('utf-8')
                return text
            except Exception as e:
                logger.warning(f"Textract failed for RTF: {e}")
        
        return f"[RTF text extraction not available - install textract]"
    
    def _extract_fallback(self, file_path: str = None, file_content: bytes = None, filename: str = "") -> str:
        """Fallback extraction for unsupported file types."""
        # Try to read as plain text
        try:
            if file_content:
                return file_content.decode('utf-8', errors='ignore')
            elif file_path:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
        except Exception as e:
            logger.warning(f"Fallback extraction failed: {e}")
        
        return f"[Text extraction not supported for {filename}]"
    
    def get_supported_extensions(self) -> list:
        """Get list of supported file extensions."""
        base_extensions = ['.txt']
        
        if self.available_extractors['pypdf2'] or self.available_extractors['pymupdf']:
            base_extensions.append('.pdf')
        
        if self.available_extractors['python_docx']:
            base_extensions.append('.docx')
        
        if self.available_extractors['textract']:
            base_extensions.extend(['.doc', '.rtf', '.odt'])
        
        return base_extensions


# Global text extractor instance
_text_extractor = None

def get_text_extractor() -> TextExtractor:
    """Get a singleton instance of TextExtractor."""
    global _text_extractor
    if _text_extractor is None:
        _text_extractor = TextExtractor()
    return _text_extractor

def extract_text_from_file(file_path: str = None, file_content: bytes = None, filename: str = None) -> str:
    """
    Convenience function to extract text from a file.
    
    Args:
        file_path: Path to the file
        file_content: File content as bytes
        filename: Original filename
    
    Returns:
        Extracted text content
    """
    extractor = get_text_extractor()
    return extractor.extract_text(file_path, file_content, filename)
